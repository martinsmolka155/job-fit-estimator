"""PDF/DOCX text extractor.

Converts uploaded CV files to raw text for downstream processing.
OCR is NOT supported in MVP — scanned PDFs return is_scanned=True with clean error message.
"""

from __future__ import annotations

import logging
from pathlib import Path

import pymupdf4llm
from docx import Document

from src.schemas import ExtractedDocument

logger = logging.getLogger(__name__)

# Below this character count we treat the PDF as scanned (no extractable text).
# Set low — pymupdf returns essentially zero chars on a true scan; legitimate
# digital CVs can be terse (50-100 chars on a one-page minimal CV).
# The gibberish detector below is the secondary filter for noisy OCR-like output.
_MIN_DIGITAL_CHARS = 30

# Maximum ratio of non-alphanumeric, non-space characters before classifying as gibberish
_GIBBERISH_THRESHOLD = 0.30

# Hard cap on PDF page count. CVs are single documents; anything over this is
# either a combined portfolio, a bulk upload, or a DoS attempt. Rejected with
# a clear ValueError that maps to HTTP 422 at the API layer.
_MAX_PDF_PAGES = 30


class UnsupportedFormatError(ValueError):
    """Raised when the input file format is not supported (not PDF or DOCX)."""


def _is_gibberish(text: str) -> bool:
    """Return True if >30% of characters are non-alphanumeric and non-space.

    Used to detect scanned PDFs where pymupdf4llm extracts garbage characters
    instead of readable text.
    """
    if not text:
        return True
    non_alnum_non_space = sum(1 for c in text if not c.isalnum() and not c.isspace())
    return (non_alnum_non_space / len(text)) > _GIBBERISH_THRESHOLD


def _extract_pdf(file_path: Path) -> ExtractedDocument:
    """Extract text from a digital PDF using pymupdf4llm.

    Returns an ``is_scanned=True`` document for:
    - corrupt or password-protected files (pymupdf raises)
    - image-only / near-empty PDFs (scanned without text layer)

    Raises:
        ValueError: If the PDF exceeds ``_MAX_PDF_PAGES`` (maps to HTTP 422).
    """
    # Check page count before expensive text extraction.  Corrupt or
    # password-protected files often surface here via a fitz exception, which
    # we convert to the scanned/unsupported sentinel so the caller can produce
    # a friendly 422 rather than a 500.
    page_count = _count_pdf_pages(file_path)
    if page_count == 0:
        # fitz failed to open the file — corrupt or encrypted PDF.
        logger.warning("PDF could not be opened (corrupt or encrypted): %s", file_path.name)
        return ExtractedDocument(
            text="",
            source_format="pdf",
            extraction_method="scanned_unsupported",
            page_count=0,
            char_count=0,
            is_scanned=True,
        )

    if page_count > _MAX_PDF_PAGES:
        raise ValueError(
            f"PDF has {page_count} pages which exceeds the {_MAX_PDF_PAGES}-page limit. "
            "Please upload a CV document, not a combined portfolio or bulk file."
        )

    try:
        # to_markdown() can return either a single str or a list of dict pages
        # depending on call signature; we always coerce to str for downstream use.
        raw = pymupdf4llm.to_markdown(str(file_path))
        text = raw if isinstance(raw, str) else "\n\n".join(p.get("text", "") for p in raw)
    except Exception:
        logger.exception("pymupdf4llm failed to extract text from PDF (logged path omitted)")
        return ExtractedDocument(
            text="",
            source_format="pdf",
            extraction_method="scanned_unsupported",
            page_count=page_count,
            char_count=0,
            is_scanned=True,
        )

    # Detect scanned PDF: too little text or mostly garbage characters
    if len(text.strip()) < _MIN_DIGITAL_CHARS or _is_gibberish(text):
        logger.warning(
            "PDF appears scanned (chars=%d). OCR not supported in MVP.",
            len(text),
        )
        return ExtractedDocument(
            text="",
            source_format="pdf",
            extraction_method="scanned_unsupported",
            page_count=page_count,
            char_count=0,
            is_scanned=True,
        )

    return ExtractedDocument(
        text=text,
        source_format="pdf",
        extraction_method="pymupdf",
        page_count=page_count,
        char_count=len(text),
        is_scanned=False,
    )


def _count_pdf_pages(file_path: Path) -> int:
    """Count pages in a PDF file.

    Returns 0 when the file cannot be opened (corrupt, encrypted, truncated).
    Callers treat 0 as the scanned/unsupported sentinel.
    Path is intentionally NOT logged to avoid leaking uploaded filenames.
    """
    try:
        import fitz  # PyMuPDF — bundled with pymupdf4llm

        with fitz.open(str(file_path)) as doc:
            return len(doc)
    except Exception:
        logger.warning("Could not count PDF pages — file may be corrupt or encrypted")
        return 0


def _extract_docx(file_path: Path) -> ExtractedDocument:
    """Extract text from a DOCX file using python-docx."""
    doc = Document(str(file_path))
    parts: list[str] = []

    # Paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    text = "\n".join(parts)
    return ExtractedDocument(
        text=text,
        source_format="docx",
        extraction_method="docx",
        page_count=1,  # DOCX has no reliable page count without rendering
        char_count=len(text),
        is_scanned=False,
    )


def extract_text(file_path: Path) -> ExtractedDocument:
    """Extract raw text from a CV file (PDF or DOCX).

    For scanned PDFs, returns ExtractedDocument with is_scanned=True and empty text.
    Callers should check is_scanned and raise a user-facing error rather than proceeding.

    Raises:
        UnsupportedFormatError: if the file extension is not .pdf or .docx
    """
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        return _extract_pdf(file_path)
    elif suffix == ".docx":
        return _extract_docx(file_path)
    else:
        raise UnsupportedFormatError(
            f"Unsupported file format: {suffix!r}. "
            "Only PDF (.pdf) and Word (.docx) files are supported."
        )
